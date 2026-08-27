# -*- coding: utf-8 -*-
"""Builds the workshop script as a plain Word document.

Deliberately plain: no headings, no tables, no header or footer, no styling
beyond the bullets the requester asked for on the questions. The uploaded
sample was three unstyled paragraphs on US Letter, and this matches it.
"""
import docx
from docx.shared import Inches, Pt

doc = docx.Document()
s = doc.sections[0]
s.page_width, s.page_height = Inches(8.5), Inches(11)
s.left_margin = s.right_margin = Inches(1)
s.top_margin = s.bottom_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(10)
normal.paragraph_format.line_spacing = 1.15

def P(t=""):
    doc.add_paragraph(t)

def LABEL(t):
    doc.add_paragraph("")
    doc.add_paragraph(t)

def B(t):
    doc.add_paragraph(t, style="List Bullet")

# ----------------------------------------------------------------- OPENING
P("Hello Team and thank you for joining this session. It is our first workshop "
  "for the utilization report.")

P("Our agenda this afternoon is to talk about the 6 dashboards. These are the "
  "dashboards from the requirements provided to us by the RAC team, the deck "
  "file sent last time.")

P("So what this session will mostly talk about is the logic and the mockup of "
  "the dashboards. We will go through them one screen at a time, and for every "
  "part of the screen I will tell you what it is, what one of them counts, and "
  "where the figure would come from. Please stop me anywhere along the way if "
  "something is not clear, or if what you are seeing is not what you had in "
  "mind when you wrote the requirement.")

P("Before we start, two things I need to set out.")

P("First, every number you will see today is illustrative. They come from our "
  "test tenant, not from ADB production. What we want you to react to is the "
  "shape of the screen, the labels, and the logic. The figures will change once "
  "we run this against the real environment.")

P("Second, there are three kinds of things on these screens, and I will tell "
  "you which one you are looking at every time. Some things are drawn and the "
  "figure behind them is real and available today. Some are drawn but they are "
  "waiting on a file or a system that we do not have yet. And some are not on "
  "the screen at all, because nothing in any system we can reach records them. "
  "That last group is the reason for a lot of the questions I will be asking "
  "you today.")

# ------------------------------------------------------------ ASSESSMENT FILE
LABEL("THE ASSESSMENT FILE")

P("Before I go through the screens, let me share our internal assessment file. "
  "This is the link to it.")

P("[paste the link here]")

P("This is where we tracked every single requirement item from your deck. It "
  "has six sheets, one for each dashboard, and 251 rows in total. Let me walk "
  "you through the columns so that after this session you can read it on your "
  "own and send us your comments directly on it.")

P("The first column is just the item number, a running count within that sheet.")

P("The second column is Requirement Items. This is the requirement itself, "
  "taken from your deck. So the wording here is yours, not ours. We kept it "
  "that way on purpose so you can recognise your own requirement.")

P("The third column is Type. This tells you what kind of element it is: a "
  "Tile, a Table column, a Graph, an Indicator, a Filter, a Behaviour, or a "
  "Panel.")

P("The fourth column is In the Mockup. This is a simple Yes or No. Yes means "
  "we have drawn it in the prototype I am about to show you. No means it is "
  "not on the screen.")

P("The fifth column is the Slide number, so you can trace any row back to the "
  "exact slide in your own deck.")

P("The sixth column is Why it is not there. This is our explanation. If we "
  "removed something, this says why we removed it. If we built it but called "
  "it something different, this says what we called it instead and the reason.")

P("The seventh column is the important one, What it needs before it can be "
  "built. For a requirement that we can produce today, this column names the "
  "table and the field we would read, and the steps to get there. For one we "
  "cannot produce yet, it names exactly what is missing, whether that is a "
  "file from RAC, a change from ITD, or a system that does not exist. So if "
  "you want to know why something is not on a screen, this is the column that "
  "answers it.")

P("Of the 251 rows, 193 are in the mockup and 54 are not.")

P("Questions on the assessment file:")
B("Is the file clear enough for your team to go through on your own after this session?")
B("Who on the RAC side will review it and send back comments?")
B("By when can we expect your feedback, so we can fold it into the next version?")

# --------------------------------------------------------------- BANK-WIDE
LABEL("DASHBOARD 1, BANK-WIDE OVERSIGHT")

P("Let us start with Bank-wide Oversight. This is the first of the six key "
  "views and it is the one your committee will most likely be looking at, "
  "because it answers the question of how much is in EDRMS and who is using "
  "it, for the whole bank, on one page.")

P("You can see from the KPIs listed at the top, there are ten tiles here, and "
  "these come straight from slide 34 of your deck. They are the following:")

B("Number of active EDRMS SharePoint sites for Department, RM, office")
B("Total number of active EDRMS SharePoint sites for Sovereign projects")
B("Total number of active EDRMS SharePoint sites for Nonsovereign projects")
B("EDRMS users with recorded activity, last 180 days")
B("Total number of documents in EDRMS")
B("Total number of records declared in EDRMS")
B("Total number of physical counterparts identified")
B("Total number of records due for disposal within 12 months")
B("Records due for retention within 12 months")
B("File plan terms in use")

P("Every one of these ten tiles is clickable, and this is your own note on "
  "slide 34. But they do not all behave the same way, and this is worth "
  "explaining because it confused us at first when we were reading the deck.")

P("The first eight tiles open a table underneath this same screen. So if I "
  "click Total number of records declared in EDRMS, the screen does not change, "
  "a table just opens below showing that same total broken down by department, "
  "office and RM. You stay on Bank-wide.")

P("The last two tiles are different. They are navigation tiles. Clicking "
  "Records due for retention takes you to the Retention and Disposal dashboard, "
  "and clicking File plan terms in use takes you to the Institutional File Plan "
  "dashboard. They carry no data of their own, they are a route to another "
  "screen. That is why in your deck slides 44 and 47 carry the Bank-wide "
  "heading while the slides after them do not. The heading marks the way in, "
  "not the ownership.")

P("Now on the labels. Two of these tiles are not worded the way you wrote them, "
  "and I want to explain why rather than have you notice it later.")

P("You asked for Total number of EDRMS users. We have written EDRMS users with "
  "recorded activity, last 180 days. The reason is that there is no list "
  "anywhere of who is entitled to use EDRMS. What Microsoft gives us is an "
  "activity report, and that report only looks back 180 days. So the number we "
  "can produce is how many people did something in SharePoint in the last 180 "
  "days. It is not a headcount of EDRMS users. If we labelled it the way you "
  "wrote it, the tile would be claiming something the data cannot support.")

P("You also asked for Total number of records due for disposal. We have written "
  "Total number of records due for disposal within 12 months, because without a "
  "window the number does not mean anything. We show it as due within 30 days, "
  "90 days and 12 months.")

P("Below the tiles is the Overview of EDRMS sites table. This is the table you "
  "drew twice, on slide 16 and again on slide 35, so we have it once. The "
  "columns are your own wording: Department, office or RM, then number of EDRMS "
  "SharePoint sites, total number of documents, total number of records "
  "declared, and total number of physical counterparts. Above it are the four "
  "indicators you asked for: sites created, sites deleted, sites archived, and "
  "sites inactive over 90 days.")

P("Two behaviours on this table. Every column heading is clickable so you can "
  "re-sort by any column. And every department name is meant to open that "
  "unit's own dashboard, which is your note on slide 35.")

P("Then the Comparison panel, which is slides 6 and 17. This sets documents "
  "against records declared for each unit, so you can see at a glance which "
  "departments are declaring and which are only storing.")

P("And last, the Records Declaration Trend. This is a cumulative curve, so "
  "each point is the running total of everything declared up to that month, "
  "not the count for that month alone. You can set the date range, and the two "
  "tiles above it show the records declared in the range you picked and the "
  "monthly average for that range.")

P("Questions on Bank-wide Oversight:")
B("Are you happy with EDRMS users with recorded activity, last 180 days as the wording, or do you want to give us a definition of an EDRMS user that we can build to?")
B("On the site table, is Department, office and RM one field, or are they three separate things we need to keep apart?")
B("Some sites carry more than one department in the system, for example CWRD and SARD on the same site. Should that site's documents count to both departments, or will RAC nominate one primary department per site?")
B("Are the three disposal windows of 30 days, 90 days and 12 months the right ones for you?")
B("Is there anything on this screen that you expected to see and cannot find?")

# ---------------------------------------------------------------- DEPARTMENT
LABEL("DASHBOARD 2, DEPARTMENT INSIGHTS")

P("The second dashboard is Department Insights. If Bank-wide is the screen for "
  "the committee, this is the screen for a department head or a records manager "
  "looking at their own unit.")

P("The whole screen is driven by the department picker at the top. You choose a "
  "department, office or RM, and everything below it refreshes for that unit "
  "only. The tiles, the tables, the charts, all of it.")

P("There are seven tiles here, from slide 53, and again all seven are "
  "clickable. They are:")

B("Total number of EDRMS compliant sites created")
B("Total number of EDRMS users with recorded activity, last 180 days")
B("Total number of site visits")
B("Total number of documents in EDRMS")
B("Total number of records declared in EDRMS")
B("Total number of physical counterparts identified")
B("Total number of records due for disposal within 12 months")

P("Each one opens its own table below, and the rows in that table are the "
  "individual sites belonging to that department, so the site rows always add "
  "back up to the tile.")

P("One label change to flag here as well. You asked for Total number of site "
  "visitors. We have written Total number of site visits. The reason is that "
  "the Microsoft report counts page views and visited pages, which are events. "
  "It does not tell us how many distinct people came. So visits we can give "
  "you, visitors we cannot.")

P("Below the tiles is the Overview of EDRMS sites section. There is a summary "
  "chart, then the site list itself with the columns you asked for: name of "
  "sites, site owners, number of documents, number of records, number of "
  "physical counterparts, and number of records due for disposal. It is sorted "
  "and paged because a single department can own more than a hundred sites.")

P("Then the Library usage section. In your deck this was six separate slides, "
  "slides 61 to 66, one for each file plan category. Rather than give you six "
  "near identical screens we put a category picker on one screen. You pick the "
  "category and the table below shows the libraries in it.")

P("I need to be honest about this section though. We can give you the number of "
  "records declared per library and the number of physical counterparts per "
  "library, because our records database knows which library a declared record "
  "sits in. What we cannot give you today is the number of documents per "
  "library, and we cannot give you the number of users per library at all. "
  "Every report we have works at site level. Nothing looks inside a site. For "
  "documents per library we would need a new scan to be built, and for users "
  "per library there is no source that will ever fill it, because Microsoft "
  "reports activity per person and per site, never per library.")

P("Finally, a few things you asked for on this screen that are not there, and I "
  "want to name them rather than let you discover them. The Go-Live date is not "
  "there, because nothing records when a site went live on EDRMS. The staff, "
  "contractor and consultant split is not there, and neither is the training "
  "completion rate, because there is no register of EDRMS users carrying "
  "employment type and no training system connected to this project. The "
  "internal and external visitor split and the access requests granted and "
  "denied are not there, because no available report records them. And "
  "everything you asked for per division is not there, because the division "
  "field exists in the system but it is empty on every single EDRMS site.")

P("Questions on Department Insights:")
B("On the library usage, is the document count per library a hard requirement, or is records declared per library enough for now?")
B("Is division something that is actually maintained by anyone today? If it is not, we should agree now to drop those columns rather than keep them open.")
B("For Go-Live date, can RAC give us a date per site, or is it acceptable to use the site creation date instead? Please note these can be years apart for a site that was converted rather than newly created.")
B("Do you have a register of EDRMS users that carries whether a person is staff, contractor or consultant?")
B("Is site visits acceptable in place of site visitors?")

# ------------------------------------------------------------------ PROJECT
LABEL("DASHBOARD 3, PROJECT INSIGHTS")

P("The third dashboard is Project Insights, and I want to set expectations "
  "before I show it. Everything on this screen is illustrative. Not one figure "
  "on it has a real source today. I am showing it to you so you can react to "
  "the shape and tell us whether it is what you wanted, and so that the two "
  "things we are missing become clear.")

P("You reach this screen either from the navigation on the left, or from tiles "
  "2 and 3 on Bank-wide, the sovereign and nonsovereign project tiles.")

P("At the top is the project list. You pick a facility type, sovereign or "
  "nonsovereign, and it lists the projects with their number and name, the "
  "number of EDRMS SharePoint sites, total documents, records declared and "
  "physical counterparts. These are your columns from slides 36 and 37. Every "
  "project row is clickable and opens that project's profile below.")

P("The profile itself is slide 38. There are eight fields: facility type, "
  "project number, project type or modality of assistance, modality number, "
  "country or economy, project status, effectivity date, and closing date.")

P("Then seven tiles, same pattern as the other screens, all clickable, each "
  "opening a table read site by site:")

B("Total number of sites created")
B("Total number of EDRMS users")
B("Total number of site visitors")
B("Total number of documents in EDRMS")
B("Total number of records declared in EDRMS")
B("Total number of physical counterparts identified")
B("Total number of records due for disposal")

P("Below that is the site list for the project, and then the three charts you "
  "drew: the user mix, the cumulative declaration trend for the project, and "
  "documents against records declared read site by site with the declaration "
  "rate alongside.")

P("Now the two things that are missing, and they are separate problems.")

P("The first one is that nothing anywhere tells us which SharePoint site "
  "belongs to which project. There is no link between the two. Until we have "
  "that, not a single figure on this screen can be produced, because we cannot "
  "even work out which sites to count.")

P("The second one is the eight profile fields at the top. Facility type, "
  "modality, country, status, effectivity date, closing date, project number "
  "and project name. Those are project attributes, and they live in an ADB "
  "project system that has never been named to us in this work by anyone. Even "
  "if you gave us the site to project list tomorrow, the top third of this "
  "screen would still be empty without that system.")

P("Questions on Project Insights:")
B("Does a list already exist anywhere that maps an EDRMS SharePoint site to a project number? If not, can RAC produce one?")
B("Which ADB system holds the project attributes, the facility type, modality, country, status and the dates? Can you give us a name and a contact?")
B("Is the sovereign and nonsovereign flag something that list would carry as well?")
B("Looking at the shape of this screen, is this what you had in mind on slide 38, or would you lay it out differently?")

# ---------------------------------------------------------------- FILE PLAN
LABEL("DASHBOARD 4, INSTITUTIONAL FILE PLAN")

P("The fourth dashboard is the Institutional File Plan. This one shows how the "
  "file plan is actually being used across the sites and libraries.")

P("At the top is the rollup, which is slide 47. It shows the total number of "
  "terms and then the count for each of your five categories: Institutional "
  "Management, Administration, People Management, Programs and Operations, and "
  "Other.")

P("Below that is the Terms by category table with your columns: heading, total "
  "terms, departments, offices and RMs provisioned, number of libraries "
  "provisioned, number of documents, number of records declared, and number of "
  "physical counterparts.")

P("Then the Terms within a category section. In your deck this was slides 48 to "
  "52, one screen per category. Same approach as the library usage, we put a "
  "picker on one screen instead of repeating five near identical screens. You "
  "pick a category and you get the top level terms in it, with the same set of "
  "columns per term, and the indicators you asked for underneath: the most used "
  "terms and the least used terms.")

P("Below that are the three remaining things every one of those slides asked "
  "for. Unused terms, which are terms with little or no usage that are "
  "candidates for review or deletion. New libraries created outside the "
  "convention. And pending requests, which covers requests for new library "
  "creation, library deletion and new terms.")

P("So the shape here is complete and it follows your deck closely. But there is "
  "one thing standing between this screen and real data, and it is the same "
  "thing for all of it. We do not have the Institutional File Plan term list. "
  "It is not in any export available to us. Every single row on this dashboard, "
  "all 24 of them in the assessment file, is waiting on that one list.")

P("I would say this is the cheapest thing to unblock in the whole report. One "
  "file from you turns a completely empty dashboard into a fully working one.")

P("Questions on the Institutional File Plan:")
B("Where is the Institutional File Plan maintained today? Is it in the SharePoint term store or somewhere else?")
B("Who owns it, and can they give us an export with one row per term and its category?")
B("Are the five categories on slide 47 still current, or has the file plan changed since that deck?")
B("For unused terms, what usage level would you consider unused? Zero, or below some number?")
B("For new libraries created outside convention, can you give us the approved naming convention so we can test against it?")

# --------------------------------------------------------- THE REMAINING TWO
LABEL("DASHBOARDS 5 AND 6, OUR PROPOSAL")

P("That brings us to the last two, Retention and Disposal, and Records and "
  "Archive Holdings. For these two I would like to put a proposal to you rather "
  "than walk through them the same way.")

P("Starting with Retention and Disposal. Half of this dashboard works today. We "
  "can show you permanent and temporary retention, by sites, by libraries and "
  "by records, and we can show you the declared records grouped by the retention "
  "label on each one. All of that comes from our records database and it is "
  "real.")

P("The half that does not work is the disposal half, and this is the important "
  "distinction. Our report can tell you that a record has reached its disposal "
  "due date. That is a calculated date and we have it. What it cannot tell you "
  "is that a disposal was actually requested, approved, declined, extended or "
  "carried out, because no system anywhere records those events. So the "
  "approver column, the approved, declined and extended statuses, the number of "
  "records disposed, the disposal completion rate, none of those can be filled. "
  "Due is not the same as disposed.")

P("And we understand that disposal is a separate release with its own system "
  "still to be established. So our proposal is that we park the disposal side "
  "for now, keep only what is due in this release, and come back to those "
  "requirements once that system exists.")

P("Then Records and Archive Holdings. This one I have to be blunt about. Of the "
  "28 requirement items on this dashboard, exactly one has a source, and that "
  "is the count of physical counterparts identified, which comes from our "
  "records database. The other 27 all need a system that records physical "
  "archive activity. The storage and retrieval requests, the boxes and folders, "
  "the locations, the requestors, the room capacity, the status of a loan or a "
  "return. None of that exists in any system we can reach, and this is not a "
  "build problem. No amount of development on our side creates data that nobody "
  "is capturing.")

P("So the same proposal for this one. We park it, and we come back to it once "
  "the system for the physical archive is established. If there is any register "
  "at all today, even a spreadsheet that someone maintains, then tell us and it "
  "changes the answer completely.")

P("To be clear on what parking means, we are not dropping these requirements. "
  "They stay in the assessment file with the reason recorded against each one, "
  "so when those systems come, we pick them up again from there rather than "
  "starting over.")

P("Questions on the remaining two:")
B("Do you agree that the disposal outcomes go with the disposal release, and that this report shows only what is due for now?")
B("Is there any existing register, system or even a spreadsheet that records RAC storage and retrieval requests today?")
B("If there is nothing, do you want Records and Archive Holdings kept in scope as a specification for later, or taken out of this report entirely?")
B("For the retention side that does work, are the retention period groupings of 3, 5, 7 and 10 years the right ones?")
B("Is there a target date for either of those systems that we should be planning around?")

# ------------------------------------------------------------------- CLOSING
LABEL("CLOSING")

P("So to summarise where we are. Four of the six dashboards we can walk through "
  "and largely build: Bank-wide Oversight, Department Insights, Project "
  "Insights and the Institutional File Plan. Of those four, two are blocked "
  "almost entirely on files that only you can give us, the site to project list "
  "and the file plan term list. The other two are mostly working today. And the "
  "last two dashboards we are proposing to park until the systems behind them "
  "exist.")

P("What we need from you after this session is the following.")

B("Your review and comments on the assessment file, with a date we can expect them by")
B("The Institutional File Plan term list, or a name of who owns it")
B("The site to project register, and the name of the ADB system that holds project attributes")
B("A decision on whether division is maintained, because eight columns depend on it")
B("A decision on the definition of an EDRMS user")
B("Confirmation that the disposal outcomes and the physical archive dashboard are parked for a later release")

P("Thank you again for your time. I will send the link to the assessment file "
  "and the prototype after this session, along with a written list of the "
  "questions we went through so you have them all in one place.")

doc.save("/home/user/Jim/EDRMS_Workshop_Script.docx")
print("saved")
