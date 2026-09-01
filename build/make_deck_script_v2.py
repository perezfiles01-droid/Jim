# -*- coding: utf-8 -*-
"""Presenter script for EDRMS_Util_Report_3.pptx.

Plain prose, spoken as written. Slide number, then what to say. No headings,
no tables, no header or footer, matching the format the requester asked for.
"""
import docx
from docx.shared import Inches, Pt

doc = docx.Document()
s = doc.sections[0]
s.page_width, s.page_height = Inches(8.5), Inches(11)
s.left_margin = s.right_margin = Inches(1)
s.top_margin = s.bottom_margin = Inches(1)
n = doc.styles["Normal"]
n.font.name = "Calibri"; n.font.size = Pt(11)
n.paragraph_format.space_after = Pt(10); n.paragraph_format.line_spacing = 1.15

def P(t=""): doc.add_paragraph(t)
def SL(t):
    doc.add_paragraph("")
    doc.add_paragraph(t)
def B(t): doc.add_paragraph(t, style="List Bullet")

# ------------------------------------------------------------------ OPENING
P("BEFORE SLIDE 1, THE TWO NEW SHEETS")

P("Good morning everyone, and thank you for joining. Before I take you through "
  "the deck, I want to spend five minutes on two new sheets we have added to "
  "the assessment file, because everything we do afterwards depends on them.")

P("[Share the file. Name the two sheets as they appear.]")

P("These two sheets are your action items. One row per requirement, for each of "
  "the two dashboards we are proposing for this release.")

P("Reading across a row, you will see the requirement in your own wording, taken "
  "from your deck. Next to it there is now a description column, which is new "
  "since you last saw this file. That column does one job: it says what the "
  "figure is actually of. Total number of documents, for example, tells you "
  "nothing on its own. The description says whether that is every document in "
  "the bank or every document in the one department on that row. The same "
  "requirement appears on both dashboards and means two different things, and "
  "that column is there so nobody has to guess which.")

P("Then our reading of it: what we think it counts, where we think the figure "
  "comes from, and whether we can build it today. Then, and this is the part we "
  "need from you, two columns for your response. A decision, which is a dropdown, "
  "and a comments box beside it.")

P("The decision column is the one we need most. Right now every row on both "
  "sheets still reads Pending Review, all one hundred and forty eight of them. "
  "Agreed, agreed with changes, deferred, needs clarification, not required. "
  "Whichever you pick tells us what to do with that row.")

P("Please do not feel you have to answer every row. Most rows just need you to "
  "confirm that the source and the maintenance are acceptable. Where we have a "
  "real question it is written in plainly, and those are the ones we need back.")

P("Two things to say about how to review it.")

P("First, apply your own test to every row, the one you set out in the last "
  "workshop. Where does this figure come from, how is it collated, and what "
  "would a records and archives team member have to do to keep it current after "
  "this project ends. If the answer to that last one is more than your team can "
  "sustain, tell us and we will take the item off rather than build something "
  "that goes stale within a year.")

P("Second, and I want to be clear about this, a blank is a useful answer. If a "
  "requirement no longer matters to you, saying so is as valuable as sourcing "
  "it. We would rather deliver a smaller report that stays correct.")

P("Thank you also for the review you have already done. Both dashboards now "
  "carry your comments against every row, and a great deal of what you will see "
  "in this deck comes straight from them. Where I quote you today, it is from "
  "that file.")

P("Questions on the two sheets before we start?")

# ------------------------------------------------------------------- SLIDES
SL("SLIDE 1")
P("So, the EDRMS Utilization Report, reconfirmation workshop for release 2026.4.")
P("What we want out of this session is one decision and a set of definitions. "
  "The decision is about scope for this release. The definitions are the words "
  "in your requirements that read more than one way, and until they are settled "
  "we cannot build the figures behind them.")

SL("SLIDE 2")
P("Here is where we are going.")
P("First the proposal on scope. Then the four data sources we actually have, and "
  "how the report is fed and kept current from them. Then the two dashboards, "
  "screen by screen, where for every item I will tell you what it counts and "
  "where it comes from. Then the terminology. Then next steps.")
P("Please stop me anywhere. If what you are seeing is not what you had in mind "
  "when you wrote the requirement, that is exactly what this session is for.")
P("One thing to set out before we start. Every number you will see today is "
  "illustrative. These are mockups, not the live report. What we want you to "
  "react to is the shape of the screen, the labels and the logic. The figures "
  "will change once we run this against the real environment.")

SL("SLIDE 3")
P("This is the proposal, and it comes from what you said in the last workshop.")
P("The suggestion was to start by focusing on the first two dashboards, and to "
  "treat those as the MVP for this release rather than delivering all six. That "
  "is what this slide sets out. Bank-wide Oversight and Department Insights in "
  "this release. The other four parked.")
P("Parked, not removed. They stay in scope. They move to a later release so "
  "that the first two can be delivered properly.")
P("On the right you can see why each of the four is parked. Project Insights "
  "waits on a register linking a SharePoint site to a project. The Institutional "
  "File Plan waits on the term list. Retention and Disposal has its disposal side "
  "waiting on a release of its own. Records and Archives Holdings waits on a "
  "system that records physical archive activity.")
P("The decision we need today is at the bottom. Option A is the two dashboard "
  "MVP. Option B is all six in one release.")
P("Let me be honest about what Option B would mean. Four of the six would arrive "
  "with most of their measures unsourced. You would be looking at screens with "
  "headings and blank columns. That is the trade, and it is your call, not ours.")
P("What we need from this slide is a consensus, so we can plan against it.")

SL("SLIDE 4")
P("These are the four sources we have today. Everything on the two MVP "
  "dashboards is built from these, and nothing else is assumed.")
P("The EDRMS database is the strongest one. Declared records, physical "
  "counterparts, retention labels, disposal due dates. EDRMS writes it as records "
  "are declared, so it is live and nobody has to maintain it.")
P("Cloud Governance gives us the EDRMS compliant site list, the department, the "
  "site owner, the status and the storage. That comes from the Workspace report.")
P("The Microsoft 365 SharePoint reports give us two things: site usage, which is "
  "file counts, page views and storage, and activity, which is per user. One "
  "important limitation there. Those reports only offer windows of 7, 30, 90 or "
  "180 days, and nothing longer. That is a product limit, not a design choice, "
  "and it is the reason anything phrased as never cannot be answered from "
  "activity alone.")
P("The fourth is the site lookup file from ITD, which maps a site to its "
  "department, division and libraries. This one is the one to watch, because it "
  "is maintained by hand. It currently covers 92 sites, of which 38 are "
  "production. Extending it across the whole estate is the single request that "
  "most increases what we can report on.")
P("Most of this was worked through with ITD and AvePoint in the technical "
  "session we held before the last workshop with you. That is where the four "
  "sources were put on the table and where we confirmed what each one can and "
  "cannot give us.")
P("[Paste the link to the sample reports here so they can see the actual columns.]")

SL("SLIDE 5")
P("This slide answers the question that was put to us twice in the last "
  "workshop: how does this data stay up to date, and who owns it once the "
  "project ends. It is a fair question and it deserves a straight answer.")
P("At the top you can see the four sources feeding in. They are joined on the "
  "site, then grouped by department or division, and that is what produces the "
  "two dashboards.")
P("The table underneath is the important part. For the EDRMS database and the "
  "Microsoft reports, the answer is nobody. They maintain themselves. No effort "
  "from your team at all.")
P("Two of them do need a person, and I want to be plain about both.")
P("Cloud Governance depends on the department and division being set correctly "
  "when a site is provisioned. If those are left blank, that site quietly drops "
  "out of every departmental figure. There is no error, no warning. The number "
  "is just wrong and looks fine.")
P("In the technical session we confirmed that the department field is already "
  "there and is mandatory at site creation, so that part is in hand. Division is "
  "the gap, and it is being backfilled.")
P("The lookup file is maintained by the BA team today, in Excel. The "
  "recommendation coming out of the technical session was to move it into the "
  "EDRMS database, with an administration page so that your team can maintain it "
  "directly rather than through a spreadsheet.")
P("That is a decision we would like your view on, because it is your team who "
  "would be using that page.")

SL("SLIDE 6")
P("Now to the first dashboard, Bank-wide Oversight.")
P("What this screen answers is: how much is in EDRMS, and who is using it, for "
  "the whole bank, on one page. This is the one your committee is most likely to "
  "look at.")
P("There are four sections. The top panel, which is ten tiles. Eight of those "
  "open a table underneath on this same screen. The last two are different, they "
  "are navigation, and clicking them takes you to another dashboard entirely. "
  "That distinction confused us when we first read your deck, and it is why some "
  "of your slides carry the Bank-wide heading while the ones after them do not. "
  "The heading marks the way in, not the ownership.")
P("Below the tiles, the overview of EDRMS sites. Every department with its sites, "
  "documents, records and physical counterparts. One row each, so they can be "
  "compared side by side.")
P("Then the comparison, which sets documents against records declared, so you "
  "can see at a glance which departments are storing but not declaring.")
P("And last the records declaration trend, which is a running total over time "
  "with a date range you choose.")
P("Again, illustrative figures. React to the shape.")

SL("SLIDE 7")
P("Now the detail on the top panel. For each item: your term, what we think it "
  "counts, where the figure comes from, and a status.")
P("The status is colour coded. Green means we can source it and build it today. "
  "Amber means the term reads more than one way, and what it means changes the "
  "figure. Red means nothing available produces it.")
P("Three of these are green and straightforward. Documents, records declared and "
  "physical counterparts all have a clean source.")
P("The amber ones are where I need you.")
P("The site count is amber, and it is amber because of a change you made "
  "yourselves. On the twenty eighth of August you narrowed this to Department HQ, "
  "and your note now says that change applies to every column that groups the "
  "same way. So this tile counts sites owned by the department HQ. Resident "
  "missions and offices do not get their own line here; they sit under division. "
  "I want that said out loud, because it changes the number.")
P("Your own review asks, and I am quoting: all records have a due date for "
  "disposal, so what does due for disposal actually mean? Due in the next 30, 60 "
  "or 90 days? You are right, and we cannot build that tile until you tell us "
  "which window you want.")
P("The second one, also yours: is this user with access to EDRMS sites? That is "
  "the definition of an EDRMS user, and it is probably the single most important "
  "word in this whole requirement set. We will come back to it on the "
  "terminology slide.")
P("The red row is the sovereign and nonsovereign project sites. Those are out of "
  "scope for this release, and that was agreed in the last workshop, so I will "
  "not dwell on it.")

SL("SLIDE 8")
P("The same treatment for the tables, the indicators and the trend.")
P("Two green here: the inactive site indicator and the declaration trend.")
P("The first row carries the same narrowing I just mentioned. Your term is still "
  "department, office or RM, because that is how it is written in your deck, but "
  "what it now counts is the owning department, narrowed to Department HQ on the "
  "twenty eighth of August. We have left your wording alone and changed only what "
  "sits underneath it, so you can see exactly what moved.")
P("Never accessed EDRMS is worth a moment. That one was ambiguous, because the "
  "activity report only looks back 180 days, so someone active 200 days ago would "
  "appear as never. In the technical session we agreed a definition that works: "
  "take the people who have access to the site, its members, owners and visitors, "
  "and check whether they have ever opened it. So this one is effectively settled, "
  "and we can build it. I would just like you to confirm that is your intent.")
P("The three red rows are the ones with no source at all today.")
P("Staff, contractors and consultants needs a register of EDRMS users carrying "
  "employment type. Nothing we can reach has that.")
P("Training completion sits with the learning team. We know who holds it, we "
  "just do not have the file yet.")
P("New documents month on month needs a document level scan, which does not "
  "exist yet. That is an ITD change rather than a reporting change.")
P("There is a fourth one that is not on this slide but is in your sheet, and I "
  "want to put it on the record rather than let it pass. Onboarded since go-live. "
  "Your latest note on it reads: not feasible for now, no metadata or identifier "
  "available to determine the go live date. We agree with that, and it is the "
  "same problem as the go-live date you will see on the next dashboard. Nothing "
  "in the system records when a department started using EDRMS. If you want it, "
  "it has to come from you as a list.")

SL("SLIDE 9")
P("Second dashboard, Department Insights. If Bank-wide is the committee screen, "
  "this is the screen a department head or a records manager uses for their own "
  "unit.")
P("The whole screen is driven by the picker at the top. You choose a department, "
  "and everything below refreshes for that department only.")
P("Four sections again. The profile, which is seven tiles each opening its own "
  "table. The overview of that department's sites. Site visits. And library usage.")
P("One thing to flag on this dashboard before we go into the detail. Your own "
  "first comment on this sheet applies to all of it: as long as department and "
  "division are backfilled. You took RM out of that list in your latest pass, "
  "which is consistent with everything else we have just said. Everything on this "
  "screen depends on that single piece of backfilling being completed.")
P("Two things have come out of this dashboard since your last pass, and I want to "
  "confirm I have read them correctly. The quick window buttons, the seven, "
  "thirty, ninety and one hundred and eighty day ones, and the month and year "
  "filter, are both gone from site visits. That is no loss, because Microsoft "
  "only offers those four windows in the first place, so we will fix on one and "
  "label the screen with it rather than pretend you can pick any range.")
P("The category picker has also come out of library usage. To be clear, the "
  "libraries are still grouped by file plan category. It is only the control that "
  "went. Tell me if that is not what you meant.")

SL("SLIDE 10")
P("The profile tiles. Same format, same colour coding.")
P("And thank you for the review here. Both sheets now carry your assessment "
  "against every row, and this table follows it. Most of it agrees with ours.")
P("So to be precise about where we are. The assessment is done, on both "
  "dashboards. What is not done is the decision. Every one of the hundred and "
  "forty eight rows still says Pending Review. That is the single most useful "
  "thing you can send back to us.")
P("The green ones are documents, records declared and physical counterparts.")
P("The amber ones are mostly the same terms we have already met. Sites created, "
  "where your review asks whether that means EDRMS sites only. EDRMS users again. "
  "Records due for disposal again.")
P("Site visitors deserves a specific note. You asked for visitors. What the "
  "Microsoft report counts is page views and visited pages, which are events, "
  "not people. So we can give you visits, but not distinct visitors. If distinct "
  "people matter to you, we need to find a different source.")
P("Go-Live date is red. Nothing anywhere records when a department went live on EDRMS. "
  "We can either take that date from you, or use the site creation date as a "
  "stand in, but I should warn you those can be years apart for a site that was "
  "converted rather than newly created.")

SL("SLIDE 11")
P("The sites, visits and libraries detail.")
P("Most of this is green. Site names, owners, documents per site, records, "
  "counterparts, page views. All sourced.")
P("Three items are not.")
P("Visitors internal and external. No available report splits a visit by whether "
  "the person is internal or external. If you know of a source, we would like to "
  "see it.")
P("Documents per library is one I want to raise directly, because your review "
  "marks it feasible and our reading is different. The lookup file gives us the "
  "library names inside each site, which is why library names are green. But "
  "nothing we have counts what is inside those libraries. Cloud Governance and "
  "the Microsoft reports both stop at the site. So either there is a source we "
  "have not seen, in which case please point us at it, or this needs the document "
  "level scan. Worth five minutes now rather than a surprise later.")
P("One amber row, access requests granted and denied. Your note asks what is "
  "meant by it, and fairly so. There is a summary page in SharePoint that holds "
  "them, but we want to confirm you mean requests to join a site, and not "
  "something wider, before we build against it.")
P("Users per library is the one I would ask you to drop. Microsoft reports "
  "activity per person and per site, never per library. There is no source today "
  "and there will not be one. Rather than leave it open, I suggest we agree now "
  "to take it off.")

SL("SLIDE 12")
P("This is the slide I would most like your time on.")
P("Each of these terms reads more than one way, and which one you mean decides "
  "what the report counts. These are not our inventions. Most of them are "
  "questions from your own review of the file.")
P("EDRMS user. It could mean anyone with access to an EDRMS site, or anyone who "
  "has used it, or anyone who has declared a record. Our proposal is the last "
  "one, because it is the only one countable today. But this is your call.")
P("Due for disposal. As you pointed out, every record has a due date, so in a "
  "sense all of them are due. Our proposal is to show it in windows: due within "
  "30 days, 90 days and 12 months.")
P("Never accessed. We have covered this. Our proposal is people with access to "
  "the site who have never opened it.")
P("Department. This one is no longer a question, and I have marked it on the "
  "slide as confirmed rather than proposed. It came out of the technical session, "
  "and then you settled it yourselves on the twenty eighth of August: Department "
  "HQ, with countries, resident missions and offices sitting underneath at "
  "division level. Your note now says that applies to every column that groups "
  "the same way, which is why several figures on this deck have changed shape "
  "since you last saw them. I am reading it back only so it is on the record.")
P("The one thing still open underneath it: should a department also be able to "
  "pull a rollup across everything beneath it, or is the HQ figure on its own "
  "what you want to see?")
P("Sites created, deleted and archived. Your review asks whether this is just "
  "EDRMS sites. Our proposal is yes, EDRMS sites only, which is what was agreed "
  "in the last workshop. The roughly three thousand other SharePoint sites stay "
  "out of this dashboard.")
P("Visitors, which we have covered. Users creating documents, where we need to "
  "know whether you mean created only or created and uploaded.")
P("And the last row, unit. Your own review asked what we meant by it, in the "
  "phrase clickable into that unit's dashboard. Given everything we have just "
  "settled, we now read unit as the department on that row, nothing smaller. "
  "I raise it because unit is also the third level of your own hierarchy, "
  "underneath division, and we are not using that level at all. Two different "
  "meanings for one word is exactly the sort of thing that produces the wrong "
  "number six months from now.")
P("If we settle these today, a large part of the report becomes buildable "
  "immediately.")

SL("SLIDE 13")
P("Next steps.")
P("First, your review of the two MVP dashboards in the two new sheets we opened "
  "with. Concretely, that means the decision column filled in on the rows that "
  "matter to you. That is the main thing we need back.")
P("Second, the terminology on the previous slide. Each one of those decides what "
  "a figure counts.")
P("Third, ITD confirming the lookup file: extending it across the estate, adding "
  "the site identifier, and agreeing where it lives and who maintains it.")
P("Fourth, we take your answers and mark every item as available, needing a "
  "definition, or having no source.")
P("And fifth, for anything left with no source, we decide together whether it is "
  "dropped or deferred. That is a decision for the team, not something we would "
  "make on our own.")
P("Can we agree a date for the review to come back to us, so we can plan around "
  "it?")

SL("SLIDE 14")
P("Thank you. To summarise what we need: a consensus on the two dashboard MVP, "
  "your review in the two new sheets, and the terminology settled.")
P("I will send the file and a written list of everything we discussed today, so "
  "you have the questions in one place rather than having to work from your "
  "notes.")
P("Thank you all for your time.")

doc.save("/home/user/Jim/EDRMS_Deck_Presenter_Script_v2.docx")
print("saved")
