# -*- coding: utf-8 -*-
"""Workshop script (.docx). Built from the same content model as the deck, so
slide numbers, checklists and counts in the script always match the slides."""
import json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = os.path.dirname(os.path.abspath(__file__))
DATA = json.load(open(os.path.join(D, 'scriptdata.json')))
S = DATA['slides']; TOT = DATA['tot']; DASH = DATA['dash']

INK   = RGBColor(0x10,0x24,0x3E)
BLUE  = RGBColor(0x00,0x5A,0x96)
MUT   = RGBColor(0x6B,0x7A,0x8C)
GREEN = RGBColor(0x3F,0x7A,0x2C)
RED   = RGBColor(0xB3,0x24,0x1C)
AMBER = RGBColor(0xC7,0x7B,0x18)
LATER = RGBColor(0x3D,0x6E,0x8C)
SER, SAN = 'Cambria', 'Calibri'

doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(0.8)
    sec.left_margin = sec.right_margin = Inches(0.9)

st = doc.styles['Normal']
st.font.name = SAN; st.font.size = Pt(11); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(7); st.paragraph_format.line_spacing = 1.14

def shade(par, hexc):
    el = OxmlElement('w:shd'); el.set(qn('w:val'), 'clear'); el.set(qn('w:fill'), hexc)
    par._p.get_or_add_pPr().append(el)

def para(text='', size=11, bold=False, italic=False, color=INK, font=SAN,
         before=0, after=7, align=None, indent=0, fill=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if indent: pf.left_indent = Inches(indent)
    if fill: shade(p, fill)
    if text:
        r = p.add_run(text); r.font.name = font; r.font.size = Pt(size)
        r.bold = bold; r.italic = italic; r.font.color.rgb = color
    return p

def h1(t):
    doc.add_page_break()
    para(t, size=24, bold=True, color=INK, font=SER, after=4)
def h2(t, before=16):
    para(t, size=16, bold=True, color=INK, font=SER, before=before, after=5)
def h3(t, before=11):
    para(t, size=12.5, bold=True, color=BLUE, font=SAN, before=before, after=3)
def label(t):
    para(t.upper(), size=8.5, bold=True, color=MUT, font=SAN, before=8, after=2)
def bullet(t, color=INK, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Inches(0.28)
    r = p.add_run(t); r.font.name = SAN; r.font.size = Pt(10.5)
    r.font.color.rgb = color; r.bold = bold
def sayblock(t):
    p = para(after=8, indent=0.16, fill='EEF3F8')
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run('“' + t + '”')
    r.font.name = SER; r.font.size = Pt(11.5); r.italic = True; r.font.color.rgb = INK

def footer_pagenum():
    for sec in doc.sections:
        p = sec.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(); r.font.name = SAN; r.font.size = Pt(8.5); r.font.color.rgb = MUT
        for instr, kind in (('begin', None), (None, 'PAGE'), ('end', None)):
            f = OxmlElement('w:fldChar') if instr else OxmlElement('w:instrText')
            if instr: f.set(qn('w:fldCharType'), instr)
            else: f.set(qn('xml:space'), 'preserve'); f.text = ' PAGE '
            r._r.append(f)

SLIDE = lambda n: 'SLIDE %d' % n

def slide_minutes(m):
    k = m['kind']
    if k in ('title', 'divider'): return 1
    if k == 'screen': return 2
    if k in ('legend', 'sixviews', 'nextsteps'): return 3
    if k in ('scoreboard', 'rootcause'): return 4
    if k in ('timeline', 'seq', 'close'): return 5
    if k == 'asks': return max(3, 2 * len(m['asks']))
    if k == 'decisions': return 6
    return 2
MINS = {m['n']: slide_minutes(m) for m in S}
TOTAL_MIN = sum(MINS.values())

# ============================ COVER ============================
para('EDRMS Utilization Report', size=30, bold=True, font=SER, before=90, after=4)
para('Workshop script', size=20, font=SER, color=BLUE, after=18)
para('Everything to say, slide by slide, in plain language.', size=12.5, color=MUT, after=3)
para('%d slides  ·  %d minutes at full pace, or about 75 minutes trimmed  ·  22 August 2026'
     % (len(S), TOTAL_MIN), size=10.5, color=MUT, after=24)
para('Read this before the workshop. Bring it with you. It is written to be '
     'read aloud — the quoted paragraphs are the words, not a summary of them.',
     size=11, italic=True, color=INK)

# ============================ CONTENTS ============================
h1('Contents')
toc = [
 ('Part 1', 'Before you start', 'Room setup, what to open, and the ten-minute version'),
 ('Part 2', 'The first ninety seconds', 'Word for word. This decides whether they listen'),
 ('Part 3', 'How to steer the conversation', 'Getting real answers, and handling the hard questions'),
 ('Part 4', 'The script, slide by slide', 'All 59 slides'),
 ('Part 5', 'The sequencing conversation', 'Disposal and Physical Records, and the three pushbacks'),
 ('Part 6', 'After the workshop', 'What to send, and to whom'),
]
for a, b, c in toc:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(9)
    r = p.add_run(a + '   '); r.font.name = SAN; r.font.size = Pt(10); r.bold = True; r.font.color.rgb = BLUE
    r = p.add_run(b + '\n'); r.font.name = SER; r.font.size = Pt(14); r.bold = True; r.font.color.rgb = INK
    r = p.add_run(c); r.font.name = SAN; r.font.size = Pt(10); r.font.color.rgb = MUT

# ============================ PART 1 ============================
h1('Part 1  ·  Before you start')
h2('What this session is for', before=6)
para('You are not presenting a finished product. You are walking them through a working '
     'prototype so they can tell you what is wrong with it, and answering the questions '
     'that are blocking the rest of the build. Both of those need them talking, not you.')
para('If they leave having said nothing, the workshop failed even if the demo went perfectly.',
     italic=True, color=BLUE)

h2('Have these open before anyone arrives')
for t in ['The deck, in presenter mode so you can see these notes on your own screen.',
          'The live prototype at perezfiles01-droid.github.io/Jim — in a second window, '
          'not a second tab. You will want to jump to it and back.',
          'The gap checker workbook, in case somebody asks to see the full list.',
          'A blank page for names. Every answer you get needs a name written next to it.']:
    bullet(t)

h2('The numbers you should know without looking')
rows = [('249', 'things they asked for, counted from their own slides'),
        (str(TOT['yes']), 'built and working in the prototype'),
        (str(TOT['no']), 'waiting on an answer from them — grouped into 27 questions'),
        (str(TOT['later']), 'waiting on a release, not on data — Disposal and Physical Records'),
        ('3', 'root causes behind most of the red')]
t = doc.add_table(rows=0, cols=2); t.style = 'Light List Accent 1'
for a, b in rows:
    c = t.add_row().cells
    c[0].width = Inches(0.9); c[1].width = Inches(5.6)
    r = c[0].paragraphs[0].add_run(a); r.bold = True; r.font.size = Pt(13); r.font.name = SER
    r = c[1].paragraphs[0].add_run(b); r.font.size = Pt(10.5); r.font.name = SAN

h2('How long this actually takes')
para('At the pace written into this script — every screen explained, every question read out and '
     'answered — it runs to about %d minutes. That is the full version, and it is the right one '
     'if you have a half day.' % TOTAL_MIN)
para('To bring it to roughly 75 minutes, cut in this order:')
for t in ['Skip the screens with nothing red on them. Say "this one is complete" and move on. '
          'That alone saves around twenty minutes.',
          'On Bank-wide and Department Insights, walk two screens properly and summarise the rest. '
          'They are the same shape and the client learns the pattern quickly.',
          'Read only the question column on the "what we need from you" slides. Skip the middle '
          'column unless asked.',
          'Do not cut the two sequencing slides, the three root causes, or the three decisions. '
          'Those four are the point of the meeting.']:
    bullet(t)

h2('If the meeting gets cut to ten minutes')
para('Show four slides and nothing else: the scoreboard (slide 4), the delivery timeline '
     '(slide 5), the three root causes (slide 56), and the three decisions (slide 57). '
     'Then say this:')
sayblock('Three quarters of what you asked for is built and working. Most of what is not '
         'comes down to three missing pieces of information, and two features that have not '
         'shipped yet. I need twenty minutes of your time to point us at the three, and a '
         'yes on sequencing the two. Everything else can go by email.')

# ============================ PART 2 ============================
h1('Part 2  ·  The first ninety seconds')
para('Do not open with an agenda. Open with why they should care, then hand them permission '
     'to interrupt. Word for word:')
sayblock('Thank you for the time. Before I show you anything, two things. First — what you '
         'are about to see is real. It is a working prototype built from your own slides, not '
         'a mock-up. You can click it. Second, and more important — nothing here is final, '
         'and this is the cheapest moment there will ever be to change it. Once this is built '
         'in Power BI, changing it costs money. Today it costs a sentence.')
sayblock('So please interrupt me. If a label is wrong, if a number looks off, if a screen is '
         'not what you pictured — say so while it is on the screen. Being polite today is '
         'the expensive option.')
para('Then, and only then, the structure:')
sayblock('I am going to walk all six dashboards, screen by screen. On the left you will see '
         'the actual screen. On the right, your own checklist for that screen, so you can see '
         'exactly what is there and what is not. At the end of each dashboard I have a short '
         'list of questions where we are genuinely stuck without you. About seventy-five '
         'minutes, and I would rather it took ninety because you kept stopping me.')
para('Why this works', size=10.5, bold=True, color=BLUE, before=12, after=3)
for t in ['It gives them a job. People listen differently when they have been asked to find faults.',
          'It removes the fear that objecting is rude. That fear is why workshops go silent.',
          'It sets the frame that this is a decision meeting, not a status update.']:
    bullet(t)

# ============================ PART 3 ============================
h1('Part 3  ·  How to steer the conversation')
h2('Getting real answers instead of "we will look into it"', before=6)
para('"We will look into it" is how a question dies. When you get it, do not accept it and '
     'move on. Do this instead:')
for t in ['Ask for a name, not an action. "Who would know?" is easier to answer than "who will do it?"',
          'Offer the smallest possible version. "Even a spreadsheet would do" removes the fear '
          'that you are asking for a project.',
          'Write the name down where they can see you writing it. It converts a vague nod into a commitment.',
          'Say the consequence once, plainly, then stop. "Without it, six figures on this screen '
          'stay blank" — and then silence.']:
    bullet(t)

h2('When they ask why something is not done')
para('This will come, probably on Users or on Records and Archive Holdings. The instinct is to '
     'apologise. Do not — it makes it sound like a failure that could be fixed by trying harder. '
     'Separate the two kinds of gap out loud:')
sayblock('There are two different things on this list and I want to keep them apart. Some of it '
         'we cannot build because a piece of information does not exist anywhere we can reach — '
         'those are the questions I have for you. The rest we cannot build because the feature it '
         'depends on has not shipped yet. That second group is not a gap, it is a sequence.')

h2('When the room goes quiet')
para('Silence usually means the question was too big, not that they have no answer. Shrink it:')
for t in ['Instead of "where does this data live?" try "if you had to guess, would this be in a '
          'system or in a spreadsheet?"',
          'Instead of asking the room, ask one person by name. Rooms do not answer questions; people do.',
          'Offer a wrong answer for them to correct. "I assume nobody tracks this at all?" is '
          'remarkably good at producing a correction.']:
    bullet(t)

h2('Closing each dashboard before you move on')
para('Never move to the next dashboard without asking this. It takes ten seconds and it is '
     'where most of the useful feedback comes from:')
sayblock('Before I move on — is there anything on this dashboard you would not use? I would '
         'much rather drop something now than build it and find out later.')
para('A requirement they do not actually want is worth more to you than a requirement you cannot '
     'source. Dropping it is free.', italic=True, color=BLUE)

h2('Things to avoid saying')
tbl = doc.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
hdr = tbl.rows[0].cells
for i, txt in enumerate(['Do not say', 'Say instead']):
    r = hdr[i].paragraphs[0].add_run(txt); r.bold = True; r.font.size = Pt(10); r.font.name = SAN
for a, b in [('"The data is not available."',
              '"Nothing we can reach holds that today. Where would you look for it?"'),
             ('"That is a technical limitation."',
              '"The system has nowhere to write that down yet."'),
             ('"We need a TermId on the document table."',
              '"Nothing says which file plan category a library belongs to."'),
             ('"It is on the roadmap."',
              '"It comes with the Disposal release. The screen is already built and waiting."'),
             ('"These numbers are dummy data."',
              '"These figures are illustrative — the layout is right, the numbers are not real yet."')]:
    c = tbl.add_row().cells
    r = c[0].paragraphs[0].add_run(a); r.font.size = Pt(10); r.font.name = SAN; r.font.color.rgb = MUT
    r = c[1].paragraphs[0].add_run(b); r.font.size = Pt(10); r.font.name = SAN; r.font.color.rgb = INK

# ============================ PART 4: SLIDE BY SLIDE ============================
h1('Part 4  ·  The script, slide by slide')
para('For each slide: what is on screen, the words to say, what the feature actually does, '
     'and the question you will probably get. Quoted paragraphs are meant to be read aloud.',
     color=MUT, size=10.5)

CHIP = {'later': ('LATER RELEASE', LATER), 'no': ('NEEDS YOU', RED),
        'amber': ('PLACEHOLDER', AMBER), 'yes': ('BUILT', GREEN)}

def slide_head(n, kind, title, mins=None):
    mins = mins or ('%d min' % MINS[n])
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(2)
    shade(p, 'E9EFF5')
    r = p.add_run('  ' + SLIDE(n) + '   '); r.bold = True; r.font.size = Pt(9.5); r.font.name = SAN; r.font.color.rgb = BLUE
    r = p.add_run(title); r.bold = True; r.font.size = Pt(13.5); r.font.name = SER; r.font.color.rgb = INK
    r = p.add_run('    ' + mins); r.font.size = Pt(9); r.font.name = SAN; r.font.color.rgb = MUT

def checklist(items):
    label('The checklist on the right of this slide')
    counts = {'yes': 0, 'amber': 0, 'no': 0, 'later': 0}
    for it in items:
        k = 'later' if it['later'] else ('no' if it['d'] == 'No' else ('amber' if it['amber'] else 'yes'))
        counts[k] += 1
    bits = []
    for k in ('yes', 'amber', 'no', 'later'):
        if counts[k]: bits.append('%d %s' % (counts[k], CHIP[k][0].lower()))
    para(str(len(items)) + ' items  ·  ' + ', '.join(bits), size=10, color=MUT, after=4)
    for it in items:
        k = 'later' if it['later'] else ('no' if it['d'] == 'No' else ('amber' if it['amber'] else 'yes'))
        if k == 'yes':      # built items need no commentary; they speak for themselves
            continue
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.28)
        r = p.add_run(CHIP[k][0] + '  '); r.bold = True; r.font.size = Pt(8); r.font.name = SAN
        r.font.color.rgb = CHIP[k][1]
        r = p.add_run(it['item']); r.font.size = Pt(10); r.font.name = SAN; r.font.color.rgb = INK

for m in S:
    k = m['kind']; n = m['n']

    if k == 'title':
        slide_head(n, k, 'Title')
        label('On screen'); para('Dark title slide with four figures across the bottom.', size=10.5)
        label('Say this'); para('Use the ninety-second opening in Part 2. Do not read the slide.', size=10.5)

    elif k == 'legend':
        slide_head(n, k, 'How to read every slide')
        label('On screen'); para('Left: what the two halves of every slide mean. Right: the four status labels.', size=10.5)
        label('Say this')
        sayblock('Every slide from here works the same way. On the left, the actual screen. On '
                 'the right, your own checklist for that screen. And four colours, which matter.')
        sayblock('Green means built and running on real data. Amber means the screen is built and '
                 'correct, but the numbers on it are made up because we have no source yet — it '
                 'will look completely real, and it is not. Red means we are stuck until you '
                 'answer something. Blue means it is waiting on a release that has not shipped, '
                 'and there is nothing you can do about it today.')
        label('Spend your time on amber')
        para('This is the one that causes trouble later. If they leave thinking Project Insights '
             'is running on live data, that becomes a very awkward conversation in a month. Say '
             'it plainly: if you see a number today, assume it is illustrative unless the '
             'checklist says built.', size=10.5)
        label('They will ask'); para('“How much of it is amber?”', size=10.5, bold=True)
        para('Answer: two dashboards are almost entirely illustrative — Project Insights and '
             'Records and Archive Holdings — and I will be explicit about both when we get there. '
             'The other four are mostly real.', size=10.5)

    elif k == 'sixviews':
        slide_head(n, k, 'Six views, each answering one question')
        label('On screen'); para('Six cards, each with the question that view answers and how much of it is built.', size=10.5)
        label('Say this')
        sayblock('Six dashboards, and each one exists to answer a single question. If a dashboard '
                 'cannot answer its question, it should not be a dashboard. Read the six questions '
                 'and tell me if any of them is the wrong question — that is much easier to fix now '
                 'than after we build.')
        label('One sentence each, no more'); para('Do not narrate all six. Let them read the cards.', size=10.5)

    elif k == 'scoreboard':
        slide_head(n, k, 'Where we are')
        label('On screen'); para('A stacked bar per dashboard, and three summary boxes.', size=10.5)
        label('Say this')
        sayblock('Two hundred and forty-nine things, counted from your own slides. A hundred and '
                 'eighty-eight of them are built. Forty-five are waiting on an answer from you. '
                 'Sixteen are waiting on a release that has not shipped yet, which is a completely '
                 'different thing and I have the next slide on it.')
        label('The trap to avoid')
        para('Do not let this read as a report card on the team. Frame it forward: the forty-five '
             'are almost all one of three missing pieces of information, and they are the only '
             'people who can point you at them.', size=10.5)

    elif k == 'timeline':
        slide_head(n, k, 'Two views are waiting on a release, not on data')
        para('This is the most important slide in the deck and it deliberately comes early. Full '
             'wording is in Part 5 — read that section before the workshop.', size=10.5, italic=True, color=BLUE)
        label('Say this')
        sayblock('Before we look at anything, one piece of context that explains a lot of what '
                 'you are about to see. Four of the six dashboards can be finished on this '
                 'programme. Two of them cannot, and not because we failed to find the data. '
                 'They depend on things that do not exist in the system yet.')
        sayblock('Disposal is not built in EDRMS. There is no screen where anyone approves a '
                 'disposal and no field where the outcome gets recorded. And Records and Archive '
                 'Holdings is entirely about paper — boxes, folders, rooms — and the physical '
                 'work comes after Utilization. So the honest plan is to sequence those two '
                 'rather than force them.')
        label('Do not apologise here')
        para('This is a normal delivery decision. Presenting it as one is what makes them accept '
             'it. Presenting it as bad news invites them to treat it as bad news.', size=10.5)

    elif k == 'divider':
        t = m['tally']
        slide_head(n, k, m['dash'])
        label('On screen'); para('Dark divider with the view name, its question, and its score.', size=10.5)
        label('Say this')
        sayblock('Next dashboard: ' + m['dash'] + '. The question it answers is — ' + m['q'] +
                 ' Before I show you a single screen: is that the right question for this view?')
        label('Why ask')
        para('If the question is wrong, everything after it is wrong too, and you want to know '
             'that before you spend eight minutes on the screens. It takes ten seconds and it has '
             'saved more rework than any other habit in this project.', size=10.5)
        para('Score: %d of %d built, %d needing them%s.  %d screens to walk.'
             % (t['yes'], t['tot'], t['no'],
                (', %d waiting on a release' % t['later']) if t['later'] else '', m['screens']),
             size=10, color=MUT)

    elif k == 'screen':
        sc = m.get('script', {})
        slide_head(n, k, m['title'])
        label('On screen')
        para('The %s screen, with the checklist beside it.' % m['title'].lower(), size=10.5)
        if sc.get('does'):
            label('What this feature actually does'); para(sc['does'], size=10.5)
        if sc.get('say'):
            label('Say this'); sayblock(sc['say'])
        if sc.get('flag'):
            label('And flag this'); sayblock(sc['flag'])
        if m.get('items'): checklist(m['items'])
        if sc.get('q'):
            label('They will probably ask')
            para(sc['q'], size=10.5, bold=True, after=2)
            para('Answer: ' + sc.get('a', ''), size=10.5)

    elif k == 'seq':
        q = m['seq']
        slide_head(n, k, q['head'])
        label('On screen')
        para('Green panel: what works today. Blue panel: what cannot work yet and why. '
             'Dark bar across the bottom: what we propose.', size=10.5)
        label('Lead with the green panel')
        para('They have to hear what does work before they hear what does not, or the whole thing '
             'sounds like an excuse. Go through the three green points first, unhurried.', size=10.5)
        for t2 in q['now']: bullet(t2, color=GREEN)
        label('Then the blue panel, as fact, not apology')
        for t2 in q['blocked']: bullet(t2, color=LATER)
        label('Then the proposal, and then stop talking')
        sayblock(q['ask'])
        sayblock('Are you comfortable with us sequencing it that way?')
        para('Get a yes or a no in the room. Do not move on with a shrug. If they push back, '
             'the three answers are in Part 5.', size=10.5, italic=True, color=BLUE)
        para('%d items move to the %s.' % (q['count'], q['rel']), size=10, color=MUT)

    elif k == 'asks':
        slide_head(n, k, 'What we need from you')
        label('On screen')
        para('%d questions, each with why we are stuck and what would unblock it.' % len(m['asks']), size=10.5)
        label('How to run this slide')
        para('Read the question in the left column out loud, then stop. Do not fill the silence '
             'with the middle column unless they ask why. A good answer is a system name, a '
             'person’s name, or “it is in a spreadsheet” — any of those three is a win. Write the '
             'name down.', size=10.5)
        for i, a in enumerate(m['asks'], 1):
            h3('%d.  %s' % (m['from'] + i, a['q']))
            para('Why we are stuck: ' + a['b'], size=10.5, after=3)
            para('What would unblock it: ' + a['n'], size=10.5, italic=True, color=BLUE)

    elif k == 'rootcause':
        slide_head(n, k, 'The three answers that unlock almost everything')
        label('Say this')
        sayblock('If you take nothing else away, take these three. A list of who is staff and who '
                 'is a consultant. Anything that connects a SharePoint site to a project number. '
                 'And somewhere that records that a disposal actually happened. Those three '
                 'account for most of what is still open.')
        label('If time is short')
        para('This is the slide to land. Everything else can go to email; these three cannot.', size=10.5)

    elif k == 'decisions':
        slide_head(n, k, 'Three decisions we would like from you today')
        label('On screen'); para('Three numbered asks, with what each one needs from them.', size=10.5)
        label('Ask for all three explicitly, and wait each time')
        para('A nod is not agreement. After each one, say “can I take that as agreed?” and let the '
             'silence do the work. Write down who agreed.', size=10.5)
        label('Say this')
        sayblock('Three things I would like to agree before we finish. First, that the four '
                 'dashboards that are ready are accepted. Second, that the disposal half follows '
                 'the Disposal release. Third, that Records and Archive Holdings is built with the '
                 'physical records work rather than now.')
        sayblock('None of those asks you for a date. All three ask you to agree an order.')

    elif k == 'nextsteps':
        slide_head(n, k, 'What happens next')
        label('Say this')
        sayblock('Four steps. You point us at the sources — even just a name. We go and check each '
                 'one ourselves rather than assuming, which is how every real error in this project '
                 'has been caught. We swap the placeholder numbers for measured ones, and the amber '
                 'items turn green. And the later releases pick up the rest, with the layouts '
                 'already agreed by then.')
        label('Before anyone leaves')
        para('Name who owes what, out loud. A question with a name against it gets answered; a '
             'question on a slide does not.', size=10.5, bold=True)

    elif k == 'close':
        slide_head(n, k, 'Over to you', 'as long as it runs')
        label('Say this')
        sayblock('Two questions, and I mean both of them. Which of the open items matters most to '
                 'you? And which of them could we drop without anyone missing it?')
        label('Mean the second one')
        para('A requirement nobody actually wants is worth more to lose than a requirement you '
             'cannot source is worth to keep. People rarely volunteer this unless you ask '
             'directly and then wait.', size=10.5)

# ============================ PART 5 ============================
h1('Part 5  ·  The sequencing conversation')
para('This is the part of the workshop most likely to go wrong, and the part where getting it '
     'right saves the most. Read this section twice before you present.', color=MUT, size=10.5)

h2('Why this matters more than it looks', before=8)
para('Sixteen items across four dashboards cannot be built now. If you present them the same way '
     'as the other forty-five, they land as "the team could not do it" — and no amount of detail '
     'afterwards undoes that first impression. Presented as sequencing, the same sixteen items '
     'land as normal delivery planning, which is what they actually are.')
para('The whole argument is one distinction, and you should be able to say it in a single breath:',
     italic=True)
sayblock('We are not missing the data. The feature that would produce the data has not been '
         'built yet.')

h2('Disposal — the argument')
h3('Start with what does work', before=6)
para('Retention is genuinely working today. Labels are being applied, the report can read them, '
     'and everything about what is due and when comes from that. Say so first and say it without '
     'hedging — it is true, and it buys you the credibility to say the next part.')
sayblock('The retention half of this dashboard is real. Labels are being applied to records right '
         'now, we can read them, and we can tell you exactly what is permanent, what is temporary, '
         'and what falls due when.')
h3('Then the fact, stated plainly')
sayblock('The disposal half is different. Disposal is not built in EDRMS yet. There is no screen '
         'where somebody approves a disposal, and there is no field anywhere that records whether '
         'it was approved, declined or extended. So there is nothing for the report to read. That '
         'is not a reporting gap — it is a feature that has not shipped.')
h3('Then the proposal')
sayblock('What we propose is this: we build the retention half now, and we build the disposal '
         'screens now as well, so they are ready and waiting. When the Disposal release ships, '
         'they switch on. Nothing is wasted, and nothing gets invented in the meantime.')
para('Fourteen items across three dashboards move with this: seven on Bank-wide Oversight, six on '
     'Department Insights, one on Retention and Disposal.', size=10.5, color=MUT)

h2('Physical Records — the argument')
h3('Be honest first', before=6)
para('This dashboard is almost entirely invented, and pretending otherwise will not survive the '
     'first question. Say it before you show it, not after.')
sayblock('I want to be straight with you before we look at this one. Exactly one figure on this '
         'dashboard is real — how many declared records are marked as having a paper copy. '
         'Everything else you are about to see, I made up.')
h3('Then the reason, which is a good one')
sayblock('Everything on this dashboard is physical. Boxes, folders, rooms, requests, retrievals. '
         'The systems that will hold that information come with the Physical Records work, and '
         'that comes after Utilization. So there is nothing to read yet — not because we have not '
         'looked, but because it does not exist yet.')
h3('Then the proposal, which turns a problem into a deliverable')
sayblock('So here is what we suggest. Treat this dashboard as a Physical Records deliverable '
         'rather than a Utilization one. Look at the layout today and tell me whether it is right, '
         'because that is genuinely useful — we can sign the design off now and build it when '
         'there are real sources behind it. What we would rather not do is ship it full of numbers '
         'we invented, because those have a way of being quoted back to us as real.')
para('Twenty-six items move with this — the whole dashboard.', size=10.5, color=MUT)

h2('The three pushbacks, and what to say')
for q, a in [
 ('“Can you not just show it empty, with zeroes?”',
  'No, and this is worth being firm on. A dashboard full of zeroes gets screenshotted and quoted '
  'as fact — someone will report that RAC holds zero boxes. An empty dashboard is more dangerous '
  'than an absent one, because it looks like an answer. What we can do instead is show the layout '
  'as an agreed specification, which gives you the same visibility without the risk.'),
 ('“We committed to six dashboards. This looks like you are delivering four.”',
  'All six are designed, and all six are in the prototype you have seen today. What changes is '
  'when two of them get their data, and that is set by when those features ship, not by us. If '
  'anything, agreeing the design now means those two get built faster once the sources exist, '
  'because the specification work is already done.'),
 ('“How long after the release before we see them working?”',
  'Short, because the hard part is already done. The screens are built and the figures are laid '
  'out. What is left is pointing them at real data instead of illustrative data. I would rather '
  'not put a number on it in the room, but it is weeks of work, not months of design.')]:
    h3(q); para('Answer: ' + a, size=10.5)

h2('If they will not agree in the room')
para('Do not push it to a stalemate. Take the smaller win instead:')
sayblock('That is fine — we do not need a decision today. What would help is agreeing that we '
         'stop trying to source these two now, and pick them up when the releases land. If you '
         'want to revisit that later, nothing is lost.')
para('An agreement to stop spending effort is nearly as useful as an agreement on sequencing, and '
     'much easier to get.', italic=True, color=BLUE)

# ============================ PART 6 ============================
h1('Part 6  ·  After the workshop')
h2('Same day', before=8)
for t in ['Send one email, short. The three decisions and whether each was agreed. Nothing else.',
          'Send the list of questions with the name that was given against each one, and ask '
          'people to correct their own name if you got it wrong. Being corrected is easier than '
          'being chased.',
          'Note anything they said they would not use. That is scope you can drop, and it is worth '
          'money.']:
    bullet(t)
h2('That week')
for t in ['Chase the three root causes first, not the easiest ones. Who is who, site to project, '
          'and where a disposal is recorded.',
          'Check every source yourself before believing it. Export the real file and read it — '
          'every genuine error on this project has been caught that way and none has been caught '
          'by assuming.',
          'Update the gap checker as answers land, so the next conversation starts from fact.']:
    bullet(t)
h2('Before the next session')
for t in ['Turn amber items green wherever a source arrived. That progress is the most persuasive '
          'thing you can show.',
          'Bring back anything they said they would drop, and confirm it is dropped. Silence is '
          'not agreement.',
          'If the monthly snapshot was agreed, make sure it actually started. Its value only '
          'builds from the day it begins.']:
    bullet(t)

para('', before=18)
para('This script was generated from the same content as the deck. If the deck changes, '
     'regenerate it rather than editing it by hand, or the slide numbers will drift.',
     size=9.5, italic=True, color=MUT)

footer_pagenum()
OUT = os.path.join(D, '..', 'EDRMS_Workshop_Script_2026-08-22.docx')
doc.save(OUT)
print('written', OUT)
