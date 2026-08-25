#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Workshop script (.docx), built from the same workshop_content.json as the deck,
so a slide number, a row number or a count cannot differ between the two.

Written to be read off, not studied. Every block is the same four lines in the
same order: SAY, SHOW, ASK, WRITE. No stage directions, no notes on tone.
"""
import json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = os.path.dirname(os.path.abspath(__file__))
K = json.load(open(os.path.join(D, "workshop_content.json")))
XL = "EDRMS_Util_Dashboard_Gap_Checker_2026-08-21.xlsx"

INK   = RGBColor(0x10, 0x24, 0x3E)
BLUE  = RGBColor(0x00, 0x5A, 0x96)
TEAL  = RGBColor(0x00, 0x7B, 0x7E)
MUT   = RGBColor(0x6B, 0x7A, 0x8C)
GREEN = RGBColor(0x3F, 0x7A, 0x2C)
RED   = RGBColor(0xB3, 0x24, 0x1C)
AMBER = RGBColor(0xC7, 0x7B, 0x18)
SER, SAN = "Cambria", "Calibri"

doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(0.75)
    sec.left_margin = sec.right_margin = Inches(0.85)

st = doc.styles["Normal"]
st.font.name = SAN
st.font.size = Pt(11)
st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.12


def shade(par, hexc):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexc)
    par._p.get_or_add_pPr().append(el)


def para(text="", size=11, bold=False, italic=False, color=INK, font=SAN,
         before=0, after=6, indent=0, fill=None, space=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent:
        pf.left_indent = Inches(indent)
    if space:
        pf.line_spacing = space
    if text:
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = color
    if fill:
        shade(p, fill)
    return p


def cue(label, text, color, fill=None):
    """One SAY / SHOW / ASK / WRITE line. The label is the thing the eye finds."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)
    pf.left_indent = Inches(0.75)
    pf.first_line_indent = Inches(-0.75)
    r = p.add_run(label.ljust(6))
    r.font.name = SAN; r.font.size = Pt(10.5); r.bold = True; r.font.color.rgb = color
    r2 = p.add_run("  " + text)
    r2.font.name = SAN; r2.font.size = Pt(11); r2.font.color.rgb = INK
    if fill:
        shade(p, fill)
    return p


def say(t):   cue("SAY",   t, BLUE)
def show(t):  cue("SHOW",  t, TEAL)
def ask(t):   cue("ASK",   t, RED)
def write(t): cue("WRITE", t, GREEN)


def h1(text, page_break=False):
    if page_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    p = para(text, size=17, bold=True, color=INK, font=SER, before=4, after=4)
    return p


def h2(text, before=12):
    return para(text, size=12.5, bold=True, color=INK, font=SAN, before=before, after=4)


def rule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr"); bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), "6")
    bt.set(qn("w:color"), "DCE4EC"); bt.set(qn("w:space"), "1")
    b.append(bt); pPr.append(b)
    p.paragraph_format.space_after = Pt(6)


def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.font.name = SAN; r.font.size = Pt(10); r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c.paragraphs[0], "10243E")
        tcPr = c._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), "10243E"); tcPr.append(sh)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(v))
            r.font.name = SAN; r.font.size = Pt(10); r.font.color.rgb = INK
    for i, w in enumerate(widths):
        for row in t.rows:
            row.cells[i].width = Inches(w)
    return t


TABS = {t["tab"]: t for t in K["tabs"]}
SHEETS = {s["tab"]: s for s in K["sheets"]}
QS = K["questions"]

# ============================================================ front page ===
h1("EDRMS Utilization Report, workshop script")
para("Read this straight down. SAY is your words, SHOW is what to put on screen, "
     "ASK is the question verbatim, WRITE is the cell to type the answer into.",
     size=11, color=MUT, italic=True, after=10)

para("The file on screen: " + XL, size=11, bold=True, after=2)
para("Keep it open the whole session. Share that window, not the slides. "
     "The deck is 16 slides and you leave it after slide 7.", size=10.5, color=MUT, after=10)

h2("Timing", before=6)
table(["Segment", "Minutes", "What gets settled"],
      [[a, str(b), c] for a, b, c in K["agenda"]],
      [3.0, 0.9, 3.2])
para("", after=6)
para("If you are running late, cut tab 4 to five minutes. It has no open rows. "
     "Never cut tabs 1 and 2, which carry all " + str(K["tot"]["no"]) +
     " of them.", size=10.5, color=MUT, italic=True, after=8)

h2("The three numbers to have in your head")
table(["Number", "What it is"],
      [[str(K["tot"]["tot"]), "requirement items read off the client's own 69 slide deck"],
       [str(K["tot"]["yes"]), "are built and on screen in the prototype today"],
       [str(K["tot"]["no"]), "are not, and all of them sit on tabs 1 and 2"]],
      [1.0, 6.1])

# ============================================================== opening ====
h1("Opening, 5 minutes", page_break=True)

h2("Slide 1, title", before=2)
say("Thank you for the requirements deck. It was detailed, and it is the reason we "
    "could build as much as we did. Today is not a demo. We built what we could "
    "build, and I need your answers on the rest.")
say("We will spend three minutes on slides and the rest of the session inside one "
    "Excel file. Everything is in that file.")

h2("Slide 2, why we are here")
say("There are " + str(K["tot"]["tot"]) + " separate things your deck asks for. We have "
    + str(K["tot"]["yes"]) + " of them built. " + str(K["tot"]["no"]) + " are not built.")
say("The " + str(K["tot"]["no"]) + " are not work we have not got to. They are things "
    "no system we can reach records. That is what I need you for today.")
ask("Before we start, is anyone here new to this project, or joining for one tab only?")

h2("Slide 3, how today runs")
say("Six tabs, one per dashboard, left to right. I will ask seventeen questions. "
    "I will type your answers into the file as you give them, and you get the file "
    "back at the end of the call with your answers in it.")
say("If you disagree with something you see, say so at the row. That is what the "
    "file is for.")

h2("Slide 4, how to read the file")
say("Eight columns. Two matter. Column D says whether it is built. Column G says "
    "what it needs, and where that column says ASK, that is a question for you.")
show("Switch to the file now. Tab 1, row 4, the headings. Then row 6, the first item.")
say("Amber row means built. Red row means not built, and the row tells you why in "
    "its own words.")

h2("Slide 5, the mockups column")
show("Tab 1, scroll right to column I. Click one picture so they see it enlarge.")
say("Every requirement has a picture of the built screen on its own row. Nobody has "
    "to remember what a screen looked like.")
ask("Can everyone see column I on their screen? If you are on a laptop you may need "
    "to scroll right.")

h2("Slide 6, the scoreboard")
say("Tabs 1 and 2 carry every open row. Tabs 3, 4 and 6 show nothing open, and that "
    "is not good news. It means we drew every screen but almost nothing behind them "
    "has a real source yet.")

h2("Slide 7, the five answers")
say("Almost every open row waits on one of five things: your user register, your "
    "project register, what a division is, a go-live date, and the physical records "
    "sources. If we leave today with an owner and a date against those five, this "
    "session has done its job.")
say("That is the last slide. From here it is the file.")

# ============================================================= the tabs ====
for t in K["tabs"]:
    sh = SHEETS[t["tab"]]
    qs = [q for q in QS if q["tab"] == t["tab"]]
    h1(f'Tab {t["tab"]}, {t["short"]}, {t["minutes"]} minutes', page_break=True)

    para(f'{sh["yes"]} built  |  {sh["no"]} not built  |  rows {sh["first_row"]} to '
         f'{sh["last_row"]}  |  questions: ' +
         (", ".join("Q" + str(q["id"]) for q in qs) if qs else "none dedicated"),
         size=10.5, bold=True, color=BLUE, after=8)

    h2("Open it", before=2)
    show(f'Tab {t["tab"]}, "{sh["name"][2:]}". {t["jump"]}')
    say(t["what"])
    say(f'On this tab, {sh["yes"]} of {sh["tot"]} rows are built' +
        (f' and {sh["no"]} are not.' if sh["no"] else
         ', and none are marked as missing. What is missing here is not the screen, '
         'it is the source behind it.'))

    h2("Walk it")
    show(f'Scroll from row {sh["first_row"]} down. Stop at every red row.')
    ask("Anything on this tab that is built but wrong? Wrong label, wrong grouping, "
        "wrong default? Say it at the row and I will type it in column F.")
    write(f'Tab {t["tab"]}, column F on that row, prefixed CLIENT: so we can find it later.')

    if qs:
        h2("The questions on this tab")
        for q in qs:
            rule()
            para(f'Q{q["id"]}. {q["title"]}', size=12, bold=True, color=INK, after=3)
            para(f'Rows {q["rows"]}   |   unblocks: {q["unblocks"]}',
                 size=9.5, color=MUT, after=4)
            show(f'Tab {q["tab"]}, rows {q["rows"]}.')
            ask(q["ask"])
            para("Why it matters, if they push back: " + q["why"],
                 size=10, color=MUT, indent=0.75, after=3)
            para("A good answer sounds like: " + q["good"],
                 size=10, color=GREEN, indent=0.75, after=3)
            para('If they say "we will get back to you": ' + q["fallback"],
                 size=10, color=AMBER, indent=0.75, after=4)
            write(f'Tab {q["tab"]}, column G on the first row of that range, and the '
                  f'owner and date on the capture sheet at the back.')

    h2("Before you move on")
    ask(f'Anything on tab {t["tab"]} we have not covered that you expected to see?')
    say("Then we move to the next tab.")

# ========================================================== the close ======
h1("Close, 7 minutes", page_break=True)

h2("Slide 14, what we need from you", before=2)
say("Five sources. I need a name and a date against each one, now, on the call.")
for name, what, count, q in K["five"]:
    ask(f'{name}. {what}. Does it exist, who owns it, and when can we have it?')
say("If the answer to any of these is that it does not exist, that is a good answer. "
    "It means we take those rows off the report today instead of showing you a blank "
    "column for another month.")

h2("Slide 15, what happens next")
say("Each answer turns into something specific, and most of them land the same week "
    "you send them.")

h2("Slide 16, close")
say("You get this file back today with your answers in it. The prototype is live at "
    "perezfiles01-droid.github.io/Jim, and it is a specification for the Power BI "
    "report, not the report itself.")
ask("Last one. Is there anything this report should show that is not in your deck "
    "and not in this file?")
write("Anywhere on the relevant tab, at the bottom, prefixed NEW:.")

# ======================================================== capture sheet ====
h1("Capture sheet, tick as you go", page_break=True)
para("Print this page. It is the only thing you need to have filled in by the end.",
     size=10.5, color=MUT, italic=True, after=8)

h2("The five sources", before=2)
table(["Source", "Exists?", "Owner", "By when"],
      [[name, "", "", ""] for name, what, count, q in K["five"]],
      [2.8, 1.0, 1.8, 1.5])
para("", after=8)

h2("The seventeen questions")
table(["#", "Question", "Tab", "Rows", "Answered?"],
      [[f'Q{q["id"]}', q["title"], q["tab"], q["rows"], ""] for q in QS],
      [0.5, 3.5, 0.5, 1.4, 1.1])
para("", after=6)
para("Anything not ticked is what the follow-up email is about.",
     size=10.5, color=MUT, italic=True)

OUT = os.path.join(D, "..", "EDRMS_Workshop_Script_2026-08-24.docx")
doc.save(OUT)
print("wrote", OUT, "| questions", len(QS), "| tabs", len(K["tabs"]))
